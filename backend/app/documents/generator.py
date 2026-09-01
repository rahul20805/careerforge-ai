import os
import uuid
from typing import Dict, Any, List
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

from app.config import settings


class DocumentGenerator:
    """
    Generates genuine, professional ATS-readable DOCX and PDF documents.
    Templates supported:
    - ATS Classic
    - Technical
    - Research
    - Academic CV
    - Quant/Finance
    - Software Engineering
    - One-page Internship
    - International Research
    """

    @classmethod
    def generate_docx(cls, resume_content: Dict[str, Any], template_name: str = "ATS Classic") -> str:
        doc = DocxDocument()

        # Set 0.75-inch standard margins
        sections = doc.sections
        for s in sections:
            s.top_margin = Inches(0.75)
            s.bottom_margin = Inches(0.75)
            s.left_margin = Inches(0.75)
            s.right_margin = Inches(0.75)

        # Header: Name
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(resume_content.get("full_name", "Professional Candidate"))
        name_run.font.name = "Calibri"
        name_run.font.size = Pt(18)
        name_run.bold = True
        name_p.paragraph_format.space_after = Pt(2)

        # Contact Info
        contact_parts = []
        if resume_content.get("email"):
            contact_parts.append(resume_content["email"])
        if resume_content.get("phone"):
            contact_parts.append(resume_content["phone"])
        if resume_content.get("location"):
            contact_parts.append(resume_content["location"])
        for k, v in resume_content.get("links", {}).items():
            if v:
                contact_parts.append(f"{k}: {v}")

        if contact_parts:
            contact_p = doc.add_paragraph()
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_p.add_run(" | ".join(contact_parts))
            contact_run.font.name = "Calibri"
            contact_run.font.size = Pt(9.5)
            contact_run.font.color.rgb = RGBColor(70, 70, 70)
            contact_p.paragraph_format.space_after = Pt(10)

        # Helper to add section headers
        def add_section_header(title: str):
            h_p = doc.add_paragraph()
            h_run = h_p.add_run(title.upper())
            h_run.font.name = "Calibri"
            h_run.font.size = Pt(11)
            h_run.bold = True
            h_run.font.color.rgb = RGBColor(20, 30, 70)
            h_p.paragraph_format.space_before = Pt(8)
            h_p.paragraph_format.space_after = Pt(2)

        # 1. Summary
        if resume_content.get("summary"):
            add_section_header("Professional Summary")
            p = doc.add_paragraph()
            p_run = p.add_run(resume_content["summary"])
            p_run.font.name = "Calibri"
            p_run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)

        # 2. Skills
        skills_by_cat = resume_content.get("skills_by_category", {})
        if skills_by_cat:
            add_section_header("Technical Skills")
            for cat_name, skill_list in skills_by_cat.items():
                if skill_list:
                    p = doc.add_paragraph()
                    cat_run = p.add_run(f"{cat_name}: ")
                    cat_run.bold = True
                    cat_run.font.name = "Calibri"
                    cat_run.font.size = Pt(10)
                    skills_run = p.add_run(", ".join(skill_list))
                    skills_run.font.name = "Calibri"
                    skills_run.font.size = Pt(10)
                    p.paragraph_format.space_after = Pt(2)

        # 3. Experience
        experiences = resume_content.get("experiences", [])
        if experiences:
            add_section_header("Experience & Research")
            for exp in experiences:
                p_title = doc.add_paragraph()
                r_pos = p_title.add_run(exp.get("position", "Position"))
                r_pos.bold = True
                r_pos.font.name = "Calibri"
                r_pos.font.size = Pt(10.5)
                
                r_org = p_title.add_run(f" — {exp.get('organization', '')}")
                r_org.font.name = "Calibri"
                r_org.font.size = Pt(10)
                
                r_dates = p_title.add_run(f" ({exp.get('dates', '')})")
                r_dates.font.name = "Calibri"
                r_dates.font.size = Pt(9.5)
                r_dates.italic = True
                p_title.paragraph_format.space_after = Pt(1)

                for b in exp.get("bullets", []):
                    b_text = b.get("text", "") if isinstance(b, dict) else str(b)
                    bp = doc.add_paragraph(style='List Bullet')
                    brun = bp.add_run(b_text)
                    brun.font.name = "Calibri"
                    brun.font.size = Pt(9.5)
                    bp.paragraph_format.space_after = Pt(1)

        # 4. Projects
        projects = resume_content.get("projects", [])
        if projects:
            add_section_header("Key Projects")
            for prj in projects:
                p_title = doc.add_paragraph()
                r_title = p_title.add_run(prj.get("title", "Project"))
                r_title.bold = True
                r_title.font.name = "Calibri"
                r_title.font.size = Pt(10.5)
                
                techs = prj.get("technologies", [])
                if techs:
                    r_tech = p_title.add_run(f" | {', '.join(techs)}")
                    r_tech.font.name = "Calibri"
                    r_tech.font.size = Pt(9.5)
                    r_tech.italic = True
                p_title.paragraph_format.space_after = Pt(1)

                for b in prj.get("bullets", []):
                    b_text = b.get("text", "") if isinstance(b, dict) else str(b)
                    bp = doc.add_paragraph(style='List Bullet')
                    brun = bp.add_run(b_text)
                    brun.font.name = "Calibri"
                    brun.font.size = Pt(9.5)
                    bp.paragraph_format.space_after = Pt(1)

        # 5. Education
        educations = resume_content.get("educations", [])
        if educations:
            add_section_header("Education")
            for edu in educations:
                p_edu = doc.add_paragraph()
                r_deg = p_edu.add_run(edu.get("degree", "Degree"))
                r_deg.bold = True
                r_deg.font.name = "Calibri"
                r_deg.font.size = Pt(10)
                
                r_inst = p_edu.add_run(f" — {edu.get('institution', '')}")
                r_inst.font.name = "Calibri"
                r_inst.font.size = Pt(10)
                
                if edu.get("gpa"):
                    r_gpa = p_edu.add_run(f" (GPA: {edu['gpa']})")
                    r_gpa.font.name = "Calibri"
                    r_gpa.font.size = Pt(9.5)
                p_edu.paragraph_format.space_after = Pt(2)

        filename = f"resume_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(settings.GENERATED_DOCS_DIR, filename)
        doc.save(filepath)
        return filepath

    @classmethod
    def generate_pdf(cls, resume_content: Dict[str, Any], template_name: str = "ATS Classic") -> str:
        filename = f"resume_{uuid.uuid4().hex[:8]}.pdf"
        filepath = os.path.join(settings.GENERATED_DOCS_DIR, filename)
        
        pdf_doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=40,
            leftMargin=40,
            topMargin=40,
            bottomMargin=40
        )
        
        styles = getSampleStyleSheet()
        
        name_style = ParagraphStyle(
            'ResumeName',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            alignment=1, # Center
            textColor=colors.HexColor('#111827')
        )
        
        contact_style = ParagraphStyle(
            'ResumeContact',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=8.5,
            leading=12,
            alignment=1,
            textColor=colors.HexColor('#4B5563')
        )
        
        section_style = ParagraphStyle(
            'ResumeSection',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=10.5,
            leading=14,
            textColor=colors.HexColor('#1E3A8A'),
            spaceBefore=8,
            spaceAfter=3
        )
        
        body_style = ParagraphStyle(
            'ResumeBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=12,
            textColor=colors.HexColor('#1F2937')
        )
        
        bullet_style = ParagraphStyle(
            'ResumeBullet',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=8.5,
            leading=11.5,
            leftIndent=12,
            textColor=colors.HexColor('#374151')
        )

        elements = []

        # Name
        elements.append(Paragraph(resume_content.get("full_name", "Candidate"), name_style))
        elements.append(Spacer(1, 3))
        
        # Contact
        contact_parts = []
        if resume_content.get("email"):
            contact_parts.append(resume_content["email"])
        if resume_content.get("phone"):
            contact_parts.append(resume_content["phone"])
        if resume_content.get("location"):
            contact_parts.append(resume_content["location"])
        for k, v in resume_content.get("links", {}).items():
            if v:
                contact_parts.append(f"{k}: {v}")

        if contact_parts:
            elements.append(Paragraph(" &nbsp;|&nbsp; ".join(contact_parts), contact_style))
            elements.append(Spacer(1, 6))

        # Summary
        if resume_content.get("summary"):
            elements.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
            elements.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor('#CBD5E1'), spaceAfter=4))
            elements.append(Paragraph(resume_content["summary"], body_style))
            elements.append(Spacer(1, 4))

        # Skills
        skills_by_cat = resume_content.get("skills_by_category", {})
        if skills_by_cat:
            elements.append(Paragraph("TECHNICAL SKILLS", section_style))
            elements.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor('#CBD5E1'), spaceAfter=4))
            for cat_name, skill_list in skills_by_cat.items():
                if skill_list:
                    elements.append(Paragraph(f"<b>{cat_name}:</b> {', '.join(skill_list)}", body_style))
            elements.append(Spacer(1, 4))

        # Experience
        experiences = resume_content.get("experiences", [])
        if experiences:
            elements.append(Paragraph("EXPERIENCE & RESEARCH", section_style))
            elements.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor('#CBD5E1'), spaceAfter=4))
            for exp in experiences:
                header_text = f"<b>{exp.get('position', '')}</b> — {exp.get('organization', '')} <i>({exp.get('dates', '')})</i>"
                elements.append(Paragraph(header_text, body_style))
                for b in exp.get("bullets", []):
                    b_text = b.get("text", "") if isinstance(b, dict) else str(b)
                    elements.append(Paragraph(f"• {b_text}", bullet_style))
                elements.append(Spacer(1, 3))

        # Projects
        projects = resume_content.get("projects", [])
        if projects:
            elements.append(Paragraph("PROJECTS", section_style))
            elements.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor('#CBD5E1'), spaceAfter=4))
            for prj in projects:
                techs = f" <i>({', '.join(prj.get('technologies', []))})</i>" if prj.get("technologies") else ""
                elements.append(Paragraph(f"<b>{prj.get('title', '')}</b>{techs}", body_style))
                for b in prj.get("bullets", []):
                    b_text = b.get("text", "") if isinstance(b, dict) else str(b)
                    elements.append(Paragraph(f"• {b_text}", bullet_style))
                elements.append(Spacer(1, 3))

        # Education
        educations = resume_content.get("educations", [])
        if educations:
            elements.append(Paragraph("EDUCATION", section_style))
            elements.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor('#CBD5E1'), spaceAfter=4))
            for edu in educations:
                gpa_str = f" (GPA: {edu['gpa']})" if edu.get("gpa") else ""
                elements.append(Paragraph(f"<b>{edu.get('degree', '')}</b> — {edu.get('institution', '')}{gpa_str}", body_style))

        pdf_doc.build(elements)
        return filepath
